"""
회의록 양식 파서 - DOCX/PDF에서 구조 추출
REQ-TMPL-002: 업로드된 양식 파일에서 섹션/필드 구조 파싱
"""

import re
from pathlib import Path

from backend.utils.logger import get_logger

logger = get_logger(__name__)


def _make_fallback(raw_text: str = "") -> dict:
    """파싱 실패 시 반환할 fallback dict 생성"""
    return {
        "sections": [],
        "fields": {},
        "has_table": False,
        "table_layout": [],
        "raw_text_preview": raw_text[:1000],
    }


class TemplateParser:
    """
    회의록 양식 파일(DOCX/PDF)에서 구조를 추출하는 파서.

    반환 형식:
        {
            "sections": [{"title": str, "level": int}, ...],
            "fields": {field_name: field_type, ...},
            "has_table": bool,
            "raw_text_preview": str  # 최대 1000자
        }
    """

    def extract_structure(self, file_path: Path, file_format: str) -> dict:
        """
        파일에서 구조 정보를 추출한다.

        Args:
            file_path: 파일 경로
            file_format: 파일 형식 ("docx" 또는 "pdf")

        Returns:
            구조 정보 dict (파싱 실패 시 fallback 반환)
        """
        file_format = file_format.lower().strip(".")

        try:
            if file_format == "docx":
                return self._parse_docx(file_path)
            elif file_format == "pdf":
                return self._parse_pdf(file_path)
            else:
                logger.warning(
                    "지원하지 않는 양식 형식",
                    file_format=file_format,
                    file_path=str(file_path),
                )
                return _make_fallback()
        except Exception as exc:
            logger.error(
                "양식 구조 추출 실패 - fallback 반환",
                error=str(exc),
                file_path=str(file_path),
                file_format=file_format,
            )
            return _make_fallback()

    def _parse_docx(self, file_path: Path) -> dict:
        """
        DOCX 파일에서 구조 추출 (python-docx 사용).

        헤딩(Heading 1~6)을 sections으로 추출하고,
        테이블 존재 여부를 감지한다.
        """
        try:
            from docx import Document
        except ImportError as exc:
            logger.error("python-docx 미설치", error=str(exc))
            return _make_fallback()

        try:
            doc = Document(str(file_path))
        except Exception as exc:
            logger.warning(
                "DOCX 파일 열기 실패",
                error=str(exc),
                file_path=str(file_path),
            )
            return _make_fallback()

        sections: list[dict] = []
        raw_lines: list[str] = []
        fields: dict = {}

        # 헤딩 및 단락 파싱
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            raw_lines.append(text)

            # 헤딩 스타일 감지 (Heading 1 ~ Heading 9)
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                sections.append({"title": text, "level": level})

        # 테이블 존재 여부 감지
        has_table = len(doc.tables) > 0

        # 테이블 내용도 raw_lines에 추가
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        raw_lines.append(cell_text)

        raw_text = "\n".join(raw_lines)

        return {
            "sections": sections,
            "fields": fields,
            "has_table": has_table,
            "raw_text_preview": raw_text[:1000],
        }

    def _parse_pdf(self, file_path: Path) -> dict:
        """
        PDF 파일에서 구조 추출 (pdfplumber 사용).

        텍스트 줄을 분석하여 섹션(번호로 시작하는 줄) 및
        테이블 존재 여부를 감지한다.
        """
        try:
            import pdfplumber
        except ImportError as exc:
            logger.error("pdfplumber 미설치", error=str(exc))
            return _make_fallback()

        try:
            with pdfplumber.open(str(file_path)) as pdf:
                sections: list[dict] = []
                raw_lines: list[str] = []
                has_table = False
                fields: dict = {}
                # 테이블에서 추출한 라벨 및 레이아웃
                table_labels: list[str] = []
                table_layout: list[dict] = []

                for page in pdf.pages:
                    # 텍스트 추출
                    page_text = page.extract_text()
                    if page_text:
                        for line in page_text.split("\n"):
                            stripped = line.strip()
                            if stripped:
                                raw_lines.append(stripped)

                    # 테이블 감지 및 레이아웃 추출
                    tables = page.extract_tables()
                    if tables:
                        has_table = True
                        for table in tables:
                            for row in table:
                                if not row:
                                    continue
                                for cell in row:
                                    if cell:
                                        raw_lines.append(str(cell).strip())

                                # 행에서 라벨 셀 위치 추출
                                row_labels: list[dict] = []
                                for col_idx, cell in enumerate(row):
                                    # 줄바꿈을 공백으로 치환하여 라벨 정규화
                                    cell_text = str(cell or "").replace("\n", " ").strip()
                                    # 라벨 조건: 한글 포함, 짧은 텍스트
                                    # 공백 구분 단어 4개 이상이면 헤더 설명으로 건너뜀
                                    # 숫자/언더스코어로 시작하면 제목 셀로 건너뜀
                                    word_count = len(cell_text.split())
                                    starts_with_special = (
                                        cell_text[0].isdigit() or cell_text[0] == '_'
                                    ) if cell_text else False
                                    if (cell_text
                                            and 1 <= len(cell_text) <= 15
                                            and word_count <= 3
                                            and not starts_with_special
                                            and any('\uAC00' <= c <= '\uD7A3' for c in cell_text)):
                                        # 라벨 다음 셀의 값 추출
                                        value = ""
                                        if col_idx + 1 < len(row) and row[col_idx + 1]:
                                            value = str(row[col_idx + 1]).replace("\n", " ").strip()
                                        row_labels.append({"label": cell_text, "value": value})

                                if not row_labels:
                                    continue

                                table_labels.extend(lab["label"] for lab in row_labels)

                                # 행 레이아웃 기록
                                if len(row_labels) == 1:
                                    table_layout.append({
                                        "type": "full",
                                        "label": row_labels[0]["label"],
                                    })
                                elif len(row_labels) >= 2:
                                    table_layout.append({
                                        "type": "split",
                                        "cells": [{"label": lab["label"]} for lab in row_labels],
                                    })

                # 섹션 추출 (2가지 방법)
                # 방법 1: 테이블 라벨 기반 (테이블이 있는 양식)
                if has_table and table_labels:
                    seen: set[str] = set()
                    for label in table_labels:
                        if label not in seen:
                            seen.add(label)
                            sections.append({"title": label, "level": 1})

                # 방법 2: 번호/헤딩 패턴 (테이블 없는 양식, fallback)
                if not sections:
                    section_pattern = re.compile(
                        r"^(?:(\d+[\.\d]*)[\.\s]+|#{1,6}\s+)(.+)$"
                    )
                    for line in raw_lines:
                        match = section_pattern.match(line)
                        if match:
                            numbering = match.group(1) or ""
                            level = len(numbering.split(".")) if numbering else 1
                            title = match.group(2).strip() if match.group(2) else line
                            sections.append({"title": title, "level": level})

                raw_text = "\n".join(raw_lines)

                return {
                    "sections": sections,
                    "fields": fields,
                    "has_table": has_table,
                    "table_layout": table_layout,
                    "raw_text_preview": raw_text[:1000],
                }

        except Exception as exc:
            logger.warning(
                "PDF 파싱 실패 - fallback 반환",
                error=str(exc),
                file_path=str(file_path),
            )
            return _make_fallback()
