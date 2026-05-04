"""
SPEC-EXPORT-002: 회의록 DOCX 생성기

python-docx 라이브러리를 사용하여 회의록 데이터를 DOCX로 변환합니다.
PDF 생성기(MinutesPDFGenerator)와 동일한 입력 포맷을 사용합니다.
"""

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches


class MinutesDOCXGenerator:
    """
    회의록 DOCX 생성기 - python-docx 기반

    회의록 데이터와 선택적 요약 데이터를 받아 DOCX bytes를 반환합니다.
    """

    # 색상 정의 (RGBColor)
    COLOR_TITLE = RGBColor(30, 30, 100)
    COLOR_SECTION = RGBColor(50, 50, 150)
    COLOR_TEXT = RGBColor(40, 40, 40)

    def generate(
        self,
        minutes_data: dict[str, Any],
        summary_data: dict[str, Any] | None = None,
    ) -> bytes:
        """
        회의록 DOCX 생성

        Args:
            minutes_data: 회의록 데이터 (segments 포함)
            summary_data: 선택적 요약 데이터

        Returns:
            DOCX 파일 바이트

        Raises:
            ValueError: minutes_data에 segments가 없거나 비어있을 때
        """
        segments = minutes_data.get("segments")
        if not segments or not isinstance(segments, list):
            raise ValueError("회의록 데이터에 유효한 segments가 없습니다.")

        doc = Document()

        # 기본 스타일 설정
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.color.rgb = self.COLOR_TEXT

        # 제목
        title = doc.add_heading("회의록", level=0)
        for run in title.runs:
            run.font.color.rgb = self.COLOR_TITLE

        # 회의 정보
        doc.add_paragraph("")
        info = doc.add_heading("회의 정보", level=1)
        for run in info.runs:
            run.font.color.rgb = self.COLOR_SECTION

        task_id = minutes_data.get("task_id", "N/A")
        doc.add_paragraph(f"Task ID: {task_id}")

        total_segments = len(segments)
        doc.add_paragraph(f"총 발화 수: {total_segments}")

        # 화자별 시간 계산
        speakers = {}
        for seg in segments:
            speaker = str(seg.get("speaker", "UNKNOWN"))
            start = float(seg.get("start", 0) or 0)
            end = float(seg.get("end", 0) or 0)
            speakers[speaker] = speakers.get(speaker, 0.0) + max(0.0, end - start)

        doc.add_paragraph(f"참여 화자 수: {len(speakers)}")
        for name, duration in sorted(speakers.items(), key=lambda x: -x[1]):
            doc.add_paragraph(f"  {name}: {duration:.1f}초", style="List Bullet")

        # 요약 섹션 (선택)
        if summary_data:
            doc.add_paragraph("")
            sum_heading = doc.add_heading("요약", level=1)
            for run in sum_heading.runs:
                run.font.color.rgb = self.COLOR_SECTION

            if isinstance(summary_data, dict):
                summary_text = summary_data.get("summary") or summary_data.get(
                    "result", ""
                )
                if summary_text:
                    doc.add_paragraph(str(summary_text))
                # 키 포인트가 있으면 추가
                key_points = summary_data.get("key_points") or summary_data.get(
                    "key_points_list", []
                )
                if key_points and isinstance(key_points, list):
                    doc.add_paragraph("")
                    doc.add_heading("주요 포인트", level=2)
                    for point in key_points:
                        doc.add_paragraph(str(point), style="List Bullet")

        # 회의 내용 (전문)
        doc.add_paragraph("")
        content_heading = doc.add_heading("회의 내용", level=1)
        for run in content_heading.runs:
            run.font.color.rgb = self.COLOR_SECTION

        for seg in segments:
            speaker = str(seg.get("speaker", "UNKNOWN"))
            start = float(seg.get("start", 0) or 0)
            text = str(seg.get("text", "")).strip()
            if not text:
                continue

            timestamp = f"[{self._format_time(start)}]"
            para = doc.add_paragraph()
            run_time = para.add_run(f"{timestamp} ")
            run_time.font.size = Pt(9)
            run_time.font.color.rgb = RGBColor(128, 128, 128)

            run_speaker = para.add_run(f"{speaker}: ")
            run_speaker.bold = True
            run_speaker.font.size = Pt(11)

            run_text = para.add_run(text)
            run_text.font.size = Pt(11)

        # BytesIO로 반환
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def _format_time(seconds: float) -> str:
        """초를 HH:MM:SS 형식으로 변환"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
