"""Document import service for searchable notes."""

from __future__ import annotations

import asyncio
import uuid
from io import BytesIO
from pathlib import Path
from typing import cast

import redis.asyncio as aioredis
from pydantic import AnyHttpUrl
from sqlalchemy.ext.asyncio import AsyncSession

from backend.schemas.external_import import (
    DocumentImportResponse,
    ExternalImportSourceType,
    ExternalTextImportRequest,
)
from backend.services.external_import_service import (
    ExternalImportService,
    ExternalImportValidationError,
)
from backend.utils.file_signature import verify_file_signature

_SUPPORTED_DOCUMENT_TYPES = {"pdf", "docx"}
_SUPPORTED_IMAGE_TYPES = {"png", "jpg", "jpeg", "webp", "heic", "heif"}
_SUPPORTED_IMPORT_TYPES = _SUPPORTED_DOCUMENT_TYPES | _SUPPORTED_IMAGE_TYPES
_MAX_DOCUMENT_SIZE_BYTES = 20 * 1024 * 1024
_MIN_EXTRACTED_TEXT_CHARS = 20


class DocumentImportService:
    """Extract user-owned document text and persist it through external import."""

    def __init__(self, external_import_service: ExternalImportService | None = None) -> None:
        self._external_import_service = external_import_service or ExternalImportService()

    async def import_document(
        self,
        *,
        filename: str,
        content: bytes,
        title: str | None,
        language: str,
        db: AsyncSession,
        redis_client: aioredis.Redis,
        owner_id: uuid.UUID | None = None,
        is_guest: bool = False,
        guest_session_id: str | None = None,
    ) -> DocumentImportResponse:
        """Import a PDF/DOCX/image document as a completed searchable minutes artifact."""
        file_type = self._file_type(filename)
        self._validate_document(filename=filename, file_type=file_type, content=content)

        extracted_text = await asyncio.to_thread(
            self._extract_text,
            file_type,
            content,
        )
        normalized_text = self._normalize_text(extracted_text)
        if len(normalized_text) < _MIN_EXTRACTED_TEXT_CHARS:
            raise ExternalImportValidationError(
                "문서에서 검색 가능한 텍스트를 충분히 추출할 수 없습니다."
            )

        document_title = (title or Path(filename).stem).strip() or filename
        source_url = self._document_source_url(filename)
        imported = await self._external_import_service.import_text(
            ExternalTextImportRequest(
                source_url=cast(AnyHttpUrl, source_url),
                title=document_title,
                content=normalized_text,
                source_type=ExternalImportSourceType.DOCUMENT,
                language=language,
            ),
            db,
            redis_client,
            owner_id=owner_id,
            is_guest=is_guest,
            guest_session_id=guest_session_id,
        )

        return DocumentImportResponse(
            **imported.model_dump(),
            file_name=filename,
            file_type=file_type,
            extracted_characters=len(normalized_text),
        )

    def _file_type(self, filename: str) -> str:
        return Path(filename or "").suffix.lstrip(".").lower()

    def _validate_document(self, *, filename: str, file_type: str, content: bytes) -> None:
        if not filename:
            raise ExternalImportValidationError("문서 파일명이 필요합니다.")
        if not content:
            raise ExternalImportValidationError("빈 문서는 가져올 수 없습니다.")
        if len(content) > _MAX_DOCUMENT_SIZE_BYTES:
            raise ExternalImportValidationError("문서 크기는 최대 20MB까지 허용됩니다.")
        if file_type not in _SUPPORTED_IMPORT_TYPES:
            raise ExternalImportValidationError("PDF, DOCX 또는 이미지 문서만 가져올 수 있습니다.")
        if not verify_file_signature(content[:16], f".{file_type}"):
            raise ExternalImportValidationError("파일 시그니처가 확장자와 일치하지 않습니다.")

    def _extract_text(self, file_type: str, content: bytes) -> str:
        if file_type == "pdf":
            return self._extract_pdf_text(content)
        if file_type == "docx":
            return self._extract_docx_text(content)
        if file_type in _SUPPORTED_IMAGE_TYPES:
            return self._extract_image_text(content)
        raise ExternalImportValidationError("지원하지 않는 문서 형식입니다.")

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            import pdfplumber
        except ImportError as exc:
            raise ExternalImportValidationError(
                "PDF 텍스트 추출 기능을 사용할 수 없습니다."
            ) from exc

        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                return "\n".join(
                    page_text for page in pdf.pages if (page_text := page.extract_text())
                )
        except Exception as exc:
            raise ExternalImportValidationError("PDF 텍스트 추출에 실패했습니다.") from exc

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ExternalImportValidationError(
                "DOCX 텍스트 추출 기능을 사용할 수 없습니다."
            ) from exc

        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise ExternalImportValidationError("DOCX 텍스트 추출에 실패했습니다.") from exc

        lines: list[str] = []
        lines.extend(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                lines.extend(cell.text for cell in row.cells)
        return "\n".join(lines)

    def _extract_image_text(self, content: bytes) -> str:
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:
            raise ExternalImportValidationError(
                "이미지 OCR 기능을 사용할 수 없습니다. Pillow와 pytesseract 런타임을 확인해주세요."
            ) from exc

        try:
            with Image.open(BytesIO(content)) as image:
                return pytesseract.image_to_string(image, lang="kor+eng")
        except Exception as exc:
            raise ExternalImportValidationError("이미지 OCR 텍스트 추출에 실패했습니다.") from exc

    def _normalize_text(self, text: str) -> str:
        lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
        return "\n".join(line for line in lines if line).strip()

    def _document_source_url(self, filename: str) -> str:
        safe_name = Path(filename).name.replace(" ", "%20")
        return f"https://local.voicetextnote/imports/documents/{safe_name}"
