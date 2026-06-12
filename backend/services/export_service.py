"""
회의록 내보내기 서비스 - PDF/DOCX 변환

SPEC-EXPORT-001: 회의록 내보내기 서비스
- 단일 회의록 내보내기
- 배치 내보내기
- 템플릿 지원
"""

import asyncio
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from dataclasses import dataclass

import fpdf
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from backend.db.models import TaskResult
from backend.schemas.export import ExportFormat, ExportFile, ExportFilter
from backend.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class ExportResult:
    """내보내기 결과"""
    task_id: str
    filename: str
    path: str
    size_bytes: int
    media_type: str
    created_at: datetime
    format: ExportFormat
    title: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None


class ExportService:
    """회의록 내보내기 서비스"""
    
    def __init__(self):
        # 템플릿 설정
        self.templates = {
            ExportFormat.pdf: self._get_pdf_templates(),
            ExportFormat.docx: self._get_docx_templates()
        }
        
        # 임시 디렉토리
        self.temp_dir = Path(tempfile.gettempdir()) / "voice_to_textnote_exports"
        self.temp_dir.mkdir(exist_ok=True)
    
    async def export_meeting(
        self,
        task_result: TaskResult,
        format: ExportFormat,
        include_summary: bool = True,
        include_action_items: bool = True,
        include_audio_analysis: bool = False
    ) -> ExportResult:
        """단일 회의록 내보내기"""
        
        # 결과 데이터 파싱
        result_data = task_result.result_data or {}
        
        # 내보내기 파일 생성
        if format == ExportFormat.pdf:
            export_file = self._create_pdf_export(
                task_result=task_result,
                result_data=result_data,
                include_summary=include_summary,
                include_action_items=include_action_items,
                include_audio_analysis=include_audio_analysis
            )
        elif format == ExportFormat.docx:
            export_file = self._create_docx_export(
                task_result=task_result,
                result_data=result_data,
                include_summary=include_summary,
                include_action_items=include_action_items,
                include_audio_analysis=include_audio_analysis
            )
        else:
            raise ValueError(f"지원되지 않는 형식: {format}")
        
        return export_file
    
    async def export_batch_meetings(
        self,
        task_results: List[TaskResult],
        format: ExportFormat,
        filters: Optional[ExportFilter] = None
    ) -> List[ExportResult]:
        """배치 회의록 내보내기"""
        
        export_files = []
        
        # 비동기 처리
        tasks = []
        for task_result in task_results:
            task = asyncio.create_task(
                self.export_meeting(
                    task_result=task_result,
                    format=format
                )
            )
            tasks.append(task)
        
        # 모든 내보내기 작업 완료 대기
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"내보내기 오류: {result}")
            else:
                export_files.append(result)
        
        return export_files
    
    def _create_pdf_export(
        self,
        task_result: TaskResult,
        result_data: Dict[str, Any],
        include_summary: bool,
        include_action_items: bool,
        include_audio_analysis: bool
    ) -> ExportResult:
        """PDF 내보내기 파일 생성"""
        
        # PDF 생성
        pdf = fpdf.FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", "", 12)
        
        # 제목
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, "회의록", ln=True, align='C')
        
        # 기본 정보
        pdf.set_font("helvetica", "", 12)
        pdf.cell(0, 10, f"작업 ID: {task_result.task_id}", ln=True)
        pdf.cell(0, 10, f"생성 시간: {task_result.created_at.strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 10, f"작업 유형: {task_result.task_type}", ln=True)
        pdf.ln(10)
        
        # 트랜스크립션 내용
        if 'transcription' in result_data:
            self._add_transcription_to_pdf(pdf, result_data['transcription'])
        
        # 요약 내용
        if include_summary and 'summary' in result_data:
            self._add_summary_to_pdf(pdf, result_data['summary'])
        
        # 액션 아이템
        if include_action_items and 'action_items' in result_data:
            self._add_action_items_to_pdf(pdf, result_data['action_items'])
        
        # 파일 저장
        filename = f"meeting_minutes_{task_result.task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = self.temp_dir / filename
        
        pdf.output(str(file_path))
        
        # 파일 정보
        size_bytes = file_path.stat().st_size
        
        return ExportResult(
            task_id=task_result.task_id,
            filename=filename,
            path=str(file_path),
            size_bytes=size_bytes,
            media_type="application/pdf",
            created_at=datetime.utcnow(),
            format=ExportFormat.pdf,
            title=task_result.task_id,
            page_count=pdf.page_no(),
            word_count=self._count_words_in_pdf(pdf)
        )
    
    def _create_docx_export(
        self,
        task_result: TaskResult,
        result_data: Dict[str, Any],
        include_summary: bool,
        include_action_items: bool,
        include_audio_analysis: bool
    ) -> ExportResult:
        """DOCX 내보내기 파일 생성"""
        
        # 문서 생성
        doc = Document()
        
        # 제목
        title = doc.add_heading('회의록', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 기본 정보
        info_para = doc.add_paragraph()
        info_para.add_run(f"작업 ID: {task_result.task_id}").bold = True
        info_para.add_run(" | ")
        info_para.add_run(f"생성 시간: {task_result.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        info_para.add_run(" | ")
        info_para.add_run(f"작업 유형: {task_result.task_type}")
        
        doc.add_paragraph()
        
        # 트랜스크립션 내용
        if 'transcription' in result_data:
            self._add_transcription_to_docx(doc, result_data['transcription'])
        
        # 요약 내용
        if include_summary and 'summary' in result_data:
            self._add_summary_to_docx(doc, result_data['summary'])
        
        # 액션 아이템
        if include_action_items and 'action_items' in result_data:
            self._add_action_items_to_docx(doc, result_data['action_items'])
        
        # 파일 저장
        filename = f"meeting_minutes_{task_result.task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = self.temp_dir / filename
        
        doc.save(str(file_path))
        
        # 파일 정보
        size_bytes = file_path.stat().st_size
        
        return ExportResult(
            task_id=task_result.task_id,
            filename=filename,
            path=str(file_path),
            size_bytes=size_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            created_at=datetime.utcnow(),
            format=ExportFormat.docx,
            title=task_result.task_id,
            page_count=len(doc.sections),
            word_count=self._count_words_in_docx(doc)
        )
    
    def _add_transcription_to_pdf(self, pdf, transcription_data):
        """PDF에 트랜스크립션 추가"""
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "1. 트랜스크립션", ln=True)
        pdf.ln(5)
        
        pdf.set_font("helvetica", "", 11)
        
        if isinstance(transcription_data, dict):
            segments = transcription_data.get('segments', [])
            for segment in segments:
                if 'text' in segment:
                    text = segment['text'].strip()
                    if text:
                        pdf.multi_cell(0, 8, text)
                        pdf.ln(2)
        elif isinstance(transcription_data, str):
            pdf.multi_cell(0, 8, transcription_data)
        
        pdf.ln(5)
    
    def _add_summary_to_pdf(self, pdf, summary_data):
        """PDF에 요약 추가"""
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "2. 요약", ln=True)
        pdf.ln(5)
        
        pdf.set_font("helvetica", "", 11)
        
        if isinstance(summary_data, dict):
            summary_text = summary_data.get('content', '') or summary_data.get('summary', '')
        elif isinstance(summary_data, str):
            summary_text = summary_data
        else:
            return
        
        if summary_text:
            pdf.multi_cell(0, 8, summary_text)
            pdf.ln(5)
    
    def _add_action_items_to_pdf(self, pdf, action_items_data):
        """PDF에 액션 아이템 추가"""
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "3. 액션 아이템", ln=True)
        pdf.ln(5)
        
        pdf.set_font("helvetica", "", 11)
        
        if isinstance(action_items_data, list):
            for i, item in enumerate(action_items_data, 1):
                if isinstance(item, dict):
                    title = item.get('title', '')
                    description = item.get('description', '')
                    assignee = item.get('assignee', '')
                    
                    if title:
                        pdf.cell(0, 8, f"{i}. {title}", ln=True)
                        if description:
                            pdf.set_font("helvetica", "I", 10)
                            pdf.multi_cell(0, 6, f"   {description}")
                            pdf.set_font("helvetica", "", 11)
                        if assignee:
                            pdf.cell(0, 6, f"   담당자: {assignee}", ln=True)
                        pdf.ln(2)
        
        pdf.ln(5)
    
    def _add_transcription_to_docx(self, doc, transcription_data):
        """DOCX에 트랜스크립션 추가"""
        doc.add_heading('1. 트랜스크립션', level=2)
        
        if isinstance(transcription_data, dict):
            segments = transcription_data.get('segments', [])
            for segment in segments:
                if 'text' in segment:
                    text = segment['text'].strip()
                    if text:
                        doc.add_paragraph(text)
        elif isinstance(transcription_data, str):
            doc.add_paragraph(transcription_data)
        
        doc.add_paragraph()
    
    def _add_summary_to_docx(self, doc, summary_data):
        """DOCX에 요약 추가"""
        doc.add_heading('2. 요약', level=2)
        
        if isinstance(summary_data, dict):
            summary_text = summary_data.get('content', '') or summary_data.get('summary', '')
        elif isinstance(summary_data, str):
            summary_text = summary_data
        else:
            return
        
        if summary_text:
            doc.add_paragraph(summary_text)
        doc.add_paragraph()
    
    def _add_action_items_to_docx(self, doc, action_items_data):
        """DOCX에 액션 아이템 추가"""
        doc.add_heading('3. 액션 아이템', level=2)
        
        if isinstance(action_items_data, list):
            for i, item in enumerate(action_items_data, 1):
                if isinstance(item, dict):
                    title = item.get('title', '')
                    description = item.get('description', '')
                    assignee = item.get('assignee', '')
                    
                    if title:
                        para = doc.add_paragraph(f"{i}. {title}")
                        
                        if description:
                            para.add_run(f"\n   {description}").italic = True
                        
                        if assignee:
                            para.add_run(f"\n   담당자: {assignee}").bold = True
                        
                        doc.add_paragraph()
    
    def _get_pdf_templates(self) -> Dict[str, Any]:
        """PDF 템플릿"""
        return {
            "default": {
                "font_family": "helvetica",
                "font_size": 12,
                "page_size": "A4",
                "margins": {
                    "top": 10,
                    "bottom": 10,
                    "left": 10,
                    "right": 10
                }
            },
            "formal": {
                "font_family": "times",
                "font_size": 11,
                "page_size": "A4",
                "margins": {
                    "top": 20,
                    "bottom": 20,
                    "left": 30,
                    "right": 30
                }
            }
        }
    
    def _get_docx_templates(self) -> Dict[str, Any]:
        """DOCX 템플릿"""
        return {
            "default": {
                "font_family": "NanumGothic",
                "font_size": 11,
                "section_spacing": 12,
                "paragraph_spacing": 6
            },
            "formal": {
                "font_family": "Batang",
                "font_size": 12,
                "section_spacing": 18,
                "paragraph_spacing": 12
            }
        }
    
    async def get_export_templates(self, format: ExportFormat) -> List[Dict[str, Any]]:
        """내보내기 템플릿 조회"""
        template_list = []
        
        for template_id, template_data in self.templates[format].items():
            template_list.append({
                "id": template_id,
                "name": template_id.capitalize(),
                "description": f"{format.value.upper()} 형식의 {template_id} 템플릿",
                "sections": ["transcription", "summary", "action_items"],
                "styling": template_data
            })
        
        return template_list
    
    def _count_words_in_pdf(self, pdf) -> int:
        """PDF 단어 수 계산 (추정)"""
        # FPDF에서 직접 단어 수 계산이 어려워 페이지 수로 추정
        return pdf.page_no() * 200  # 평균 페이지당 200단어 추정
    
    def _count_words_in_docx(self, doc) -> int:
        """DOCX 단어 수 계산"""
        word_count = 0
        for paragraph in doc.paragraphs:
            word_count += len(paragraph.text.split())
        return word_count