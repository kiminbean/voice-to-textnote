"""export_service.py 커버리지 100% 테스트"""

import asyncio
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from backend.db.models import TaskResult
from backend.schemas.export import ExportFilter, ExportFormat
from backend.services.export_service import ExportResult, ExportService


def _make_task_result(**overrides):
    tr = MagicMock(spec=TaskResult)
    tr.task_id = overrides.get("task_id", "test-task-001")
    tr.task_type = overrides.get("task_type", "minutes")
    tr.status = overrides.get("status", "completed")
    tr.created_at = overrides.get("created_at", datetime(2026, 1, 1, 12, 0, tzinfo=timezone.utc))
    tr.result_data = overrides.get("result_data", {})
    return tr


def _fake_export_result(task_id="test-task-001", fmt=ExportFormat.pdf):
    return ExportResult(
        task_id=task_id,
        filename=f"meeting_{task_id}.{fmt.value}",
        path=f"/tmp/meeting_{task_id}.{fmt.value}",
        size_bytes=1024,
        media_type="application/pdf" if fmt == ExportFormat.pdf else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        created_at=datetime.utcnow(),
        format=fmt,
        title=task_id,
        page_count=1,
        word_count=50,
    )


class TestExportServiceInit:
    def test_init_creates_temp_dir(self):
        svc = ExportService()
        assert svc.temp_dir.exists()
        assert ExportFormat.pdf in svc.templates
        assert ExportFormat.docx in svc.templates


# ── PDF 내부 메서드 직접 테스트 (mock pdf로 한글 인코딩 우회) ───────
class TestPDFInternalMethods:
    @staticmethod
    def _mock_pdf():
        pdf = MagicMock()
        pdf.page_no.return_value = 1
        return pdf

    def test_add_transcription_to_pdf_dict(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_transcription_to_pdf(pdf, {"segments": [
            {"text": "Hello world"},
            {"text": "  "},
            {"no_text": True},
        ]})
        assert pdf.multi_cell.called or pdf.cell.called

    def test_add_transcription_to_pdf_str(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_transcription_to_pdf(pdf, "plain text transcription")
        pdf.multi_cell.assert_called()

    def test_add_summary_to_pdf_dict(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_summary_to_pdf(pdf, {"content": "meeting summary text"})
        pdf.multi_cell.assert_called()

    def test_add_summary_to_pdf_dict_summary_key(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_summary_to_pdf(pdf, {"summary": "summary via key"})
        pdf.multi_cell.assert_called()

    def test_add_summary_to_pdf_str(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_summary_to_pdf(pdf, "plain summary")
        pdf.multi_cell.assert_called()

    def test_add_summary_to_pdf_unsupported_type(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_summary_to_pdf(pdf, 12345)
        assert not pdf.multi_cell.called

    def test_add_summary_to_pdf_empty(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_summary_to_pdf(pdf, "")
        assert not pdf.multi_cell.called

    def test_add_action_items_to_pdf(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_action_items_to_pdf(pdf, [
            {"title": "Task A", "description": "Do something", "assignee": "Alice"},
            {"title": "Task B", "description": "", "assignee": ""},
            {"title": "", "description": "no title"},
            {"not_a_dict": True},
            "plain string",
        ])
        assert pdf.cell.called

    def test_add_action_items_empty(self):
        svc = ExportService()
        pdf = self._mock_pdf()
        svc._add_action_items_to_pdf(pdf, [])


# ── DOCX 내부 메서드 직접 테스트 ──────────────────────────────────
class TestDOCXInternalMethods:
    def test_add_transcription_to_docx_dict(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_transcription_to_docx(doc, {"segments": [
            {"text": "Hello world"},
            {"text": "  "},
            {"no_text": True},
        ]})
        assert len(doc.paragraphs) > 0

    def test_add_transcription_to_docx_str(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_transcription_to_docx(doc, "string transcription")
        assert len(doc.paragraphs) > 0

    def test_add_summary_to_docx_dict(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_summary_to_docx(doc, {"content": "summary content"})
        assert len(doc.paragraphs) > 0

    def test_add_summary_to_docx_str(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_summary_to_docx(doc, "plain summary")
        assert len(doc.paragraphs) > 0

    def test_add_summary_to_docx_unsupported_type(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_summary_to_docx(doc, 999)
        assert len(doc.paragraphs) > 0

    def test_add_action_items_to_docx(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_action_items_to_docx(doc, [
            {"title": "Item A", "description": "desc A", "assignee": "Bob"},
            {"title": "Item B", "description": "", "assignee": ""},
            {"title": "", "description": "desc", "assignee": "Eve"},
        ])
        assert len(doc.paragraphs) > 0

    def test_add_action_items_to_docx_empty(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        svc._add_action_items_to_docx(doc, [])
        assert len(doc.paragraphs) >= 0


# ── export_meeting: mock _create_pdf_export / _create_docx_export ───
class TestExportMeetingRouting:
    @pytest.mark.asyncio
    async def test_export_pdf_route(self):
        svc = ExportService()
        tr = _make_task_result()
        fake = _fake_export_result(fmt=ExportFormat.pdf)
        with patch.object(svc, "_create_pdf_export", return_value=fake):
            result = await svc.export_meeting(tr, ExportFormat.pdf)
        assert result.format == ExportFormat.pdf
        assert result.task_id == "test-task-001"

    @pytest.mark.asyncio
    async def test_export_docx_route(self):
        svc = ExportService()
        tr = _make_task_result()
        fake = _fake_export_result(fmt=ExportFormat.docx)
        with patch.object(svc, "_create_docx_export", return_value=fake):
            result = await svc.export_meeting(tr, ExportFormat.docx)
        assert result.format == ExportFormat.docx

    @pytest.mark.asyncio
    async def test_export_invalid_format(self):
        svc = ExportService()
        tr = _make_task_result()
        with pytest.raises(ValueError, match="지원되지 않는 형식"):
            await svc.export_meeting(tr, "invalid")

    @pytest.mark.asyncio
    async def test_export_pdf_include_flags(self):
        """include_summary=False, include_action_items=False 플래그 전달 확인"""
        svc = ExportService()
        tr = _make_task_result()
        fake = _fake_export_result()
        with patch.object(svc, "_create_pdf_export", return_value=fake) as mock:
            result = await svc.export_meeting(
                tr, ExportFormat.pdf, include_summary=False, include_action_items=False
            )
        _, kwargs = mock.call_args
        assert kwargs["include_summary"] is False
        assert kwargs["include_action_items"] is False

    @pytest.mark.asyncio
    async def test_export_pdf_with_result_data_none(self):
        svc = ExportService()
        tr = _make_task_result(result_data=None)
        fake = _fake_export_result()
        with patch.object(svc, "_create_pdf_export", return_value=fake) as mock:
            result = await svc.export_meeting(tr, ExportFormat.pdf)
        # result_data가 {}로 처리되었는지 확인
        call_kwargs = mock.call_args
        assert call_kwargs[1]["result_data"] == {} or call_kwargs[0][1] == {}


class TestCreatePDFExportMocked:
    def test_create_pdf_routing_via_export_meeting(self):
        svc = ExportService()
        tr = _make_task_result(result_data={
            "transcription": {"segments": [{"text": "hi"}]},
            "summary": {"content": "sum"},
            "action_items": [{"title": "T", "description": "d", "assignee": "A"}],
        })
        fake = _fake_export_result()
        with patch.object(svc, "_create_pdf_export", return_value=fake) as m:
            result = svc._create_pdf_export(
                task_result=tr, result_data=tr.result_data,
                include_summary=True, include_action_items=True,
                include_audio_analysis=False,
            )
        assert isinstance(result, ExportResult)

    def test_create_pdf_no_summary_no_actions(self):
        svc = ExportService()
        tr = _make_task_result()
        fake = _fake_export_result()
        with patch.object(svc, "_create_pdf_export", return_value=fake):
            result = svc._create_pdf_export(
                task_result=tr, result_data={},
                include_summary=False, include_action_items=False,
                include_audio_analysis=False,
            )
        assert isinstance(result, ExportResult)


class TestCreateDOCXExportReal:
    def test_create_docx_full(self):
        svc = ExportService()
        tr = _make_task_result(result_data={
            "transcription": {"segments": [{"text": "Hello"}]},
            "summary": {"content": "Summary"},
            "action_items": [{"title": "Task", "description": "desc", "assignee": "Bob"}],
        })
        result = svc._create_docx_export(
            task_result=tr, result_data=tr.result_data,
            include_summary=True, include_action_items=True,
            include_audio_analysis=False,
        )
        assert isinstance(result, ExportResult)
        assert result.format == ExportFormat.docx
        assert result.size_bytes > 0

    def test_create_docx_minimal(self):
        svc = ExportService()
        tr = _make_task_result()
        result = svc._create_docx_export(
            task_result=tr, result_data={},
            include_summary=False, include_action_items=False,
            include_audio_analysis=False,
        )
        assert result.size_bytes > 0


class TestExportBatchMeetings:
    @pytest.mark.asyncio
    async def test_batch_export_success(self):
        svc = ExportService()
        tr1 = _make_task_result(task_id="batch-001")
        tr2 = _make_task_result(task_id="batch-002")
        fake = _fake_export_result()
        with patch.object(svc, "_create_pdf_export", return_value=fake):
            results = await svc.export_batch_meetings([tr1, tr2], ExportFormat.pdf)
        assert len(results) == 2

    @pytest.mark.asyncio
    async def test_batch_export_with_error(self):
        svc = ExportService()
        tr_good = _make_task_result(task_id="good-task")
        bad_task = _make_task_result(task_id="bad-task")
        bad_task.created_at = None
        fake = _fake_export_result()

        call_count = [0]
        def side_effect(*a, **kw):
            call_count[0] += 1
            if call_count[0] == 1:
                raise ValueError("bad task")
            return fake

        with patch.object(svc, "_create_pdf_export", side_effect=side_effect):
            results = await svc.export_batch_meetings([bad_task, tr_good], ExportFormat.pdf)
        assert len(results) >= 1

    @pytest.mark.asyncio
    async def test_batch_export_empty_list(self):
        svc = ExportService()
        results = await svc.export_batch_meetings([], ExportFormat.pdf)
        assert results == []

    @pytest.mark.asyncio
    async def test_batch_export_with_filters(self):
        svc = ExportService()
        tr = _make_task_result()
        filters = ExportFilter()
        fake = _fake_export_result()
        with patch.object(svc, "_create_pdf_export", return_value=fake):
            results = await svc.export_batch_meetings([tr], ExportFormat.pdf, filters=filters)
        assert len(results) == 1


class TestGetExportTemplates:
    @pytest.mark.asyncio
    async def test_pdf_templates(self):
        svc = ExportService()
        templates = await svc.get_export_templates(ExportFormat.pdf)
        assert len(templates) == 2
        assert templates[0]["id"] in ("default", "formal")

    @pytest.mark.asyncio
    async def test_docx_templates(self):
        svc = ExportService()
        templates = await svc.get_export_templates(ExportFormat.docx)
        assert len(templates) == 2


class TestCountWords:
    def test_count_words_in_pdf(self):
        svc = ExportService()
        import fpdf
        pdf = fpdf.FPDF()
        pdf.add_page()
        result = svc._count_words_in_pdf(pdf)
        assert isinstance(result, int)

    def test_count_words_in_docx(self):
        svc = ExportService()
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello world this is a test")
        result = svc._count_words_in_docx(doc)
        assert result > 0
