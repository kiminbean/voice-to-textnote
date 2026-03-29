"""
TemplateParser 단위 테스트 (RED phase)
REQ-TMPL-002: DOCX/PDF 파일에서 구조 추출
목표: 90% 이상 커버리지
"""

from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# 테스트용 픽스처 - 샘플 파일 생성
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """
    테스트용 DOCX 파일 생성 (python-docx 사용)
    """
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("회의록", level=1)
        doc.add_heading("1. 회의 개요", level=2)
        doc.add_paragraph("일시: 2025-01-01")
        doc.add_paragraph("참석자: 홍길동, 김철수")
        doc.add_heading("2. 주요 안건", level=2)
        doc.add_paragraph("안건 1: 프로젝트 진행 현황")
        doc.add_heading("3. 결정 사항", level=2)
        doc.add_paragraph("결정 1: 다음 주 까지 보고서 제출")
        doc.add_heading("4. 다음 단계", level=2)
        doc.add_paragraph("담당자: 홍길동")

        # 테이블 추가
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "항목"
        table.rows[0].cells[1].text = "내용"
        table.rows[1].cells[0].text = "담당자"
        table.rows[1].cells[1].text = "홍길동"

        docx_path = tmp_path / "sample.docx"
        doc.save(str(docx_path))
        return docx_path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    """
    테스트용 PDF 파일 생성 (reportlab 또는 fpdf2 사용)
    reportlab이 없으면 최소한의 PDF 바이트로 대체
    """
    pdf_path = tmp_path / "sample.pdf"
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "회의록")
        c.drawString(100, 730, "1. 회의 개요")
        c.drawString(100, 710, "일시: 2025-01-01")
        c.drawString(100, 690, "2. 주요 안건")
        c.drawString(100, 670, "안건 1: 프로젝트 진행 현황")
        c.save()
    except ImportError:
        try:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            pdf.cell(200, 10, "Meeting Minutes", ln=True)
            pdf.cell(200, 10, "1. Overview", ln=True)
            pdf.cell(200, 10, "Date: 2025-01-01", ln=True)
            pdf.cell(200, 10, "2. Main Agenda", ln=True)
            pdf.output(str(pdf_path))
        except ImportError:
            # 최소 PDF 바이트 (텍스트 없음 - 파싱 실패 경로 테스트용)
            pdf_path.write_bytes(
                b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\nendobj\n"
                b"xref\n0 2\n0000000000 65535 f\n0000000009 00000 n\n"
                b"trailer\n<< /Size 2 /Root 1 0 R >>\nstartxref\n49\n%%EOF"
            )
    return pdf_path


@pytest.fixture
def empty_docx_path(tmp_path: Path) -> Path:
    """내용이 없는 DOCX 파일"""
    try:
        from docx import Document

        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(str(docx_path))
        return docx_path
    except ImportError:
        pytest.skip("python-docx not installed")


# ---------------------------------------------------------------------------
# TemplateParser 임포트 테스트
# ---------------------------------------------------------------------------


class TestTemplateParserImport:
    """TemplateParser 클래스 임포트 및 기본 구조 테스트"""

    def test_template_parser_can_be_imported(self):
        """TemplateParser 클래스를 임포트할 수 있어야 함"""
        from backend.pipeline.template_parser import TemplateParser

        assert TemplateParser is not None

    def test_template_parser_has_extract_structure_method(self):
        """extract_structure 메서드가 존재해야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        assert hasattr(parser, "extract_structure")
        assert callable(parser.extract_structure)

    def test_template_parser_has_parse_docx_method(self):
        """_parse_docx 메서드가 존재해야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        assert hasattr(parser, "_parse_docx")

    def test_template_parser_has_parse_pdf_method(self):
        """_parse_pdf 메서드가 존재해야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        assert hasattr(parser, "_parse_pdf")


# ---------------------------------------------------------------------------
# extract_structure() 반환값 구조 테스트
# ---------------------------------------------------------------------------


class TestExtractStructureReturnFormat:
    """extract_structure() 반환값 형식 검증"""

    def test_extract_structure_returns_dict(self, sample_docx_path):
        """extract_structure()는 dict를 반환해야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert isinstance(result, dict)

    def test_extract_structure_has_sections_key(self, sample_docx_path):
        """반환값에 'sections' 키 존재"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert "sections" in result

    def test_extract_structure_has_fields_key(self, sample_docx_path):
        """반환값에 'fields' 키 존재"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert "fields" in result

    def test_extract_structure_has_has_table_key(self, sample_docx_path):
        """반환값에 'has_table' 키 존재"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert "has_table" in result

    def test_extract_structure_has_raw_text_preview_key(self, sample_docx_path):
        """반환값에 'raw_text_preview' 키 존재"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert "raw_text_preview" in result

    def test_sections_is_list(self, sample_docx_path):
        """sections은 리스트여야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert isinstance(result["sections"], list)

    def test_fields_is_dict(self, sample_docx_path):
        """fields는 dict여야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert isinstance(result["fields"], dict)

    def test_has_table_is_bool(self, sample_docx_path):
        """has_table은 bool이어야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert isinstance(result["has_table"], bool)

    def test_raw_text_preview_is_str(self, sample_docx_path):
        """raw_text_preview는 str이어야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert isinstance(result["raw_text_preview"], str)


# ---------------------------------------------------------------------------
# DOCX 파싱 테스트
# ---------------------------------------------------------------------------


class TestParseDocx:
    """DOCX 파일 파싱 테스트"""

    def test_docx_extracts_headings_as_sections(self, sample_docx_path):
        """DOCX의 heading이 sections로 추출되어야 함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        # 헤딩이 있으므로 sections가 비어있으면 안 됨
        assert len(result["sections"]) > 0

    def test_docx_detects_table(self, sample_docx_path):
        """테이블이 있는 DOCX → has_table=True"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert result["has_table"] is True

    def test_docx_raw_text_preview_not_empty(self, sample_docx_path):
        """DOCX 내용 있으면 raw_text_preview 비어있으면 안 됨"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert len(result["raw_text_preview"]) > 0

    def test_docx_raw_text_preview_max_1000_chars(self, sample_docx_path):
        """raw_text_preview는 최대 1000자"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        assert len(result["raw_text_preview"]) <= 1000

    def test_empty_docx_returns_empty_sections(self, empty_docx_path):
        """내용 없는 DOCX → sections 빈 리스트"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(empty_docx_path, "docx")
        assert isinstance(result["sections"], list)

    def test_empty_docx_has_table_false(self, empty_docx_path):
        """내용 없는 DOCX → has_table=False"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(empty_docx_path, "docx")
        assert result["has_table"] is False


# ---------------------------------------------------------------------------
# PDF 파싱 테스트
# ---------------------------------------------------------------------------


class TestParsePdf:
    """PDF 파일 파싱 테스트"""

    def test_pdf_extract_structure_returns_dict(self, sample_pdf_path):
        """PDF → dict 반환"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_pdf_path, "pdf")
        assert isinstance(result, dict)

    def test_pdf_has_required_keys(self, sample_pdf_path):
        """PDF 결과에 필수 키 포함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_pdf_path, "pdf")
        assert "sections" in result
        assert "fields" in result
        assert "has_table" in result
        assert "raw_text_preview" in result

    def test_pdf_raw_text_preview_max_1000_chars(self, sample_pdf_path):
        """PDF raw_text_preview 최대 1000자"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_pdf_path, "pdf")
        assert len(result["raw_text_preview"]) <= 1000


# ---------------------------------------------------------------------------
# 실패 케이스 - fallback 동작 테스트
# ---------------------------------------------------------------------------


class TestExtractStructureFallback:
    """파싱 실패 시 fallback 동작 테스트"""

    def test_nonexistent_file_returns_fallback(self, tmp_path):
        """존재하지 않는 파일 → fallback dict 반환 (예외 없음)"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        fake_path = tmp_path / "nonexistent.docx"
        # 예외 없이 fallback 반환
        result = parser.extract_structure(fake_path, "docx")
        assert isinstance(result, dict)
        assert result["sections"] == []
        assert result["fields"] == {}
        assert result["has_table"] is False

    def test_unsupported_format_returns_fallback(self, tmp_path):
        """지원하지 않는 형식 → fallback dict 반환"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        fake_path = tmp_path / "test.txt"
        fake_path.write_text("테스트 내용")
        result = parser.extract_structure(fake_path, "txt")
        assert isinstance(result, dict)
        assert result["sections"] == []
        assert result["fields"] == {}

    def test_fallback_result_has_all_keys(self, tmp_path):
        """fallback 결과에도 모든 필수 키 포함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        fake_path = tmp_path / "nonexistent.pdf"
        result = parser.extract_structure(fake_path, "pdf")
        assert "sections" in result
        assert "fields" in result
        assert "has_table" in result
        assert "raw_text_preview" in result

    def test_fallback_raw_text_preview_is_string(self, tmp_path):
        """fallback raw_text_preview는 문자열"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        fake_path = tmp_path / "nonexistent.docx"
        result = parser.extract_structure(fake_path, "docx")
        assert isinstance(result["raw_text_preview"], str)

    def test_parse_docx_exception_returns_fallback(self, tmp_path):
        """_parse_docx 내부 예외 → fallback 반환"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        # 잘못된 DOCX (텍스트 파일)
        bad_docx = tmp_path / "bad.docx"
        bad_docx.write_bytes(b"This is not a valid DOCX file")
        result = parser.extract_structure(bad_docx, "docx")
        assert isinstance(result, dict)
        assert result["sections"] == []

    def test_parse_pdf_exception_returns_fallback(self, tmp_path):
        """_parse_pdf 내부 예외 → fallback 반환 (예외 없음)"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        # 잘못된 PDF
        bad_pdf = tmp_path / "bad.pdf"
        bad_pdf.write_bytes(b"This is not a valid PDF file")
        result = parser.extract_structure(bad_pdf, "pdf")
        assert isinstance(result, dict)
        assert result["sections"] == []


# ---------------------------------------------------------------------------
# sections 형식 테스트
# ---------------------------------------------------------------------------


class TestSectionsFormat:
    """sections 리스트의 각 항목 형식 검증"""

    def test_section_items_have_title(self, sample_docx_path):
        """sections의 각 항목에 'title' 키 포함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        for section in result["sections"]:
            assert "title" in section

    def test_section_items_have_level(self, sample_docx_path):
        """sections의 각 항목에 'level' 키 포함"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        for section in result["sections"]:
            assert "level" in section

    def test_section_title_is_string(self, sample_docx_path):
        """section의 title은 문자열"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        for section in result["sections"]:
            assert isinstance(section["title"], str)

    def test_section_level_is_int(self, sample_docx_path):
        """section의 level은 정수"""
        from backend.pipeline.template_parser import TemplateParser

        parser = TemplateParser()
        result = parser.extract_structure(sample_docx_path, "docx")
        for section in result["sections"]:
            assert isinstance(section["level"], int)
