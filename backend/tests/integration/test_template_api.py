"""
Template API 통합 테스트 (RED phase)
REQ-TMPL-001: POST /api/v1/templates - 양식 파일 업로드
REQ-TMPL-003: GET /api/v1/templates, GET /api/v1/templates/{id}, DELETE /api/v1/templates/{id}
"""

import json
import uuid
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi.testclient import TestClient

# ---------------------------------------------------------------------------
# 테스트용 DOCX 바이트 생성 헬퍼
# ---------------------------------------------------------------------------


def _make_docx_bytes() -> bytes:
    """테스트용 DOCX 바이트 생성"""
    from docx import Document

    doc = Document()
    doc.add_heading("회의록 양식", level=1)
    doc.add_heading("1. 회의 개요", level=2)
    doc.add_paragraph("일시: ")
    doc.add_paragraph("장소: ")
    doc.add_heading("2. 주요 안건", level=2)
    doc.add_heading("3. 결정 사항", level=2)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes() -> bytes:
    """테스트용 PDF 바이트 생성"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.drawString(100, 750, "Meeting Minutes Template")
        c.drawString(100, 730, "1. Overview")
        c.drawString(100, 710, "2. Agenda")
        c.save()
        return buf.getvalue()
    except ImportError:
        # 최소 PDF 바이트
        return (
            b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\nendobj\n"
            b"xref\n0 2\n0000000000 65535 f\n0000000009 00000 n\n"
            b"trailer\n<< /Size 2 /Root 1 0 R >>\nstartxref\n49\n%%EOF"
        )


# ---------------------------------------------------------------------------
# 테스트용 Redis 및 앱 픽스처
# ---------------------------------------------------------------------------


@pytest.fixture
def mock_tmpl_redis_client():
    """Redis 비동기 클라이언트 mock (Template API 전용)"""
    redis_mock = AsyncMock()
    redis_mock.get.return_value = None
    redis_mock.set.return_value = True
    redis_mock.setex.return_value = True
    redis_mock.delete.return_value = 1
    redis_mock.ping.return_value = True
    redis_mock.scard.return_value = 0
    redis_mock.keys.return_value = []
    return redis_mock


@pytest.fixture
def tmpl_client(mock_tmpl_redis_client, tmp_path):
    """
    Template API TestClient
    - Redis mock
    - 파일 저장소: tmp_path
    - WhisperEngine, DiarizationEngine mock
    """
    from backend.app.config import Settings
    from backend.app.dependencies import get_redis_client
    from backend.app.main import app
    from backend.app.middleware.auth import verify_api_key

    # 테스트용 Settings mock
    test_settings = MagicMock(spec=Settings)
    test_settings.max_concurrent_summaries = 2
    test_settings.summary_result_ttl = 86400
    test_settings.max_concurrent_minutes = 3
    test_settings.minutes_result_ttl = 86400
    test_settings.max_concurrent_diarizations = 2
    test_settings.diarization_result_ttl = 86400
    test_settings.temp_dir = tmp_path / "temp"
    test_settings.results_dir = tmp_path / "results"
    test_settings.templates_dir = tmp_path / "templates"
    test_settings.huggingface_token = "hf_testtoken"
    test_settings.diarization_model = "pyannote/speaker-diarization-3.1"
    test_settings.cache_ttl_seconds = 604800
    test_settings.temp_dir.mkdir(parents=True, exist_ok=True)
    test_settings.results_dir.mkdir(parents=True, exist_ok=True)
    test_settings.templates_dir.mkdir(parents=True, exist_ok=True)

    async def override_redis():
        return mock_tmpl_redis_client

    app.dependency_overrides[get_redis_client] = override_redis

    async def override_verify_api_key():
        return "test-bypass"

    app.dependency_overrides[verify_api_key] = override_verify_api_key

    with patch("backend.app.main.WhisperEngine") as mock_whisper_cls:
        mock_whisper_inst = MagicMock()
        mock_whisper_inst.is_loaded = True
        mock_whisper_inst.load.return_value = None
        mock_whisper_cls.get_instance.return_value = mock_whisper_inst

        with patch("backend.app.main.DiarizationEngine") as mock_dia_cls:
            mock_dia_inst = MagicMock()
            mock_dia_inst.is_loaded = True
            mock_dia_inst.load.return_value = None
            mock_dia_cls.get_instance.return_value = mock_dia_inst

            with patch("backend.app.api.v1.admin.templates.settings", test_settings):
                yield TestClient(app, raise_server_exceptions=True)

    app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# POST /api/v1/templates 테스트 (REQ-TMPL-001)
# ---------------------------------------------------------------------------


class TestPostTemplates:
    """POST /api/v1/templates - 양식 업로드"""

    def test_upload_docx_returns_201(self, tmpl_client):
        """DOCX 업로드 → 201 Created"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "테스트 회의록 양식"},
        )
        assert response.status_code == 201

    def test_upload_pdf_returns_201(self, tmpl_client):
        """PDF 업로드 → 201 Created"""
        pdf_bytes = _make_pdf_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={"file": ("template.pdf", pdf_bytes, "application/pdf")},
            data={"name": "PDF 회의록 양식"},
        )
        assert response.status_code == 201

    def test_upload_returns_template_id(self, tmpl_client):
        """업로드 응답에 template_id 포함"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "테스트 양식"},
        )
        data = response.json()
        assert "template_id" in data

    def test_upload_returns_name(self, tmpl_client):
        """업로드 응답에 name 포함"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "My Template"},
        )
        data = response.json()
        assert data["name"] == "My Template"

    def test_upload_returns_format(self, tmpl_client):
        """업로드 응답에 format 포함"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "테스트"},
        )
        data = response.json()
        assert "format" in data
        assert data["format"] == "docx"

    def test_upload_returns_structure(self, tmpl_client):
        """업로드 응답에 structure 포함"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "테스트"},
        )
        data = response.json()
        assert "structure" in data
        assert isinstance(data["structure"], dict)

    def test_upload_returns_created_at(self, tmpl_client):
        """업로드 응답에 created_at 포함"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "테스트"},
        )
        data = response.json()
        assert "created_at" in data

    def test_upload_unsupported_format_returns_422(self, tmpl_client):
        """지원하지 않는 파일 형식 → 422 Unprocessable Entity"""
        response = tmpl_client.post(
            "/api/v1/templates",
            files={"file": ("template.txt", b"text content", "text/plain")},
            data={"name": "텍스트 파일"},
        )
        assert response.status_code == 422

    def test_upload_file_too_large_returns_422(self, tmpl_client):
        """10MB 초과 파일 → 422"""
        # 10MB + 1byte
        large_content = b"A" * (10 * 1024 * 1024 + 1)
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "big.docx",
                    large_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "큰 파일"},
        )
        assert response.status_code == 422

    def test_upload_without_name_uses_filename(self, tmpl_client):
        """name 없이 업로드 → 파일명을 name으로 사용"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "my_template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        # 400이나 422가 아니라 성공해야 함
        assert response.status_code == 201

    def test_upload_saves_file_to_storage(self, tmpl_client, tmp_path):
        """업로드 후 파일이 storage에 저장되어야 함"""
        docx_bytes = _make_docx_bytes()
        response = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "저장 테스트"},
        )
        assert response.status_code == 201
        data = response.json()
        template_id = data["template_id"]
        # storage/templates/{template_id}/original.docx 파일이 존재해야 함
        expected_path = tmp_path / "templates" / template_id / "original.docx"
        assert expected_path.exists()


# ---------------------------------------------------------------------------
# GET /api/v1/templates 테스트 (REQ-TMPL-003)
# ---------------------------------------------------------------------------


class TestGetTemplateList:
    """GET /api/v1/templates - 목록 조회"""

    def test_get_list_returns_200(self, tmpl_client):
        """템플릿 목록 조회 → 200"""
        response = tmpl_client.get("/api/v1/templates")
        assert response.status_code == 200

    def test_get_list_returns_array(self, tmpl_client, mock_tmpl_redis_client):
        """목록 응답은 배열"""
        mock_tmpl_redis_client.keys.return_value = []
        response = tmpl_client.get("/api/v1/templates")
        data = response.json()
        assert isinstance(data, list)

    def test_get_list_after_upload_contains_item(self, tmpl_client, mock_tmpl_redis_client):
        """업로드 후 목록에 항목 포함"""
        # 먼저 업로드
        docx_bytes = _make_docx_bytes()
        upload_resp = tmpl_client.post(
            "/api/v1/templates",
            files={
                "file": (
                    "template.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"name": "목록 테스트"},
        )
        assert upload_resp.status_code == 201
        template_id = upload_resp.json()["template_id"]

        # Redis mock에 해당 키 설정
        meta = {
            "template_id": template_id,
            "name": "목록 테스트",
            "format": "docx",
            "created_at": "2025-01-01T00:00:00",
        }
        mock_tmpl_redis_client.keys.return_value = [f"template:{template_id}"]
        mock_tmpl_redis_client.get.return_value = json.dumps(meta)

        response = tmpl_client.get("/api/v1/templates")
        data = response.json()
        assert len(data) >= 1
        ids = [item["template_id"] for item in data]
        assert template_id in ids

    def test_list_item_has_required_fields(self, tmpl_client, mock_tmpl_redis_client):
        """목록 항목에 template_id, name, format, created_at 포함"""
        template_id = str(uuid.uuid4())
        meta = {
            "template_id": template_id,
            "name": "테스트 양식",
            "format": "docx",
            "created_at": "2025-01-01T00:00:00",
        }
        mock_tmpl_redis_client.keys.return_value = [f"template:{template_id}"]
        mock_tmpl_redis_client.get.return_value = json.dumps(meta)

        response = tmpl_client.get("/api/v1/templates")
        data = response.json()
        assert len(data) == 1
        item = data[0]
        assert "template_id" in item
        assert "name" in item
        assert "format" in item
        assert "created_at" in item


# ---------------------------------------------------------------------------
# GET /api/v1/templates/{template_id} 테스트 (REQ-TMPL-003)
# ---------------------------------------------------------------------------


class TestGetTemplateDetail:
    """GET /api/v1/templates/{template_id} - 상세 조회"""

    def test_get_detail_returns_200(self, tmpl_client, mock_tmpl_redis_client):
        """존재하는 템플릿 상세 조회 → 200"""
        template_id = str(uuid.uuid4())
        meta = {
            "template_id": template_id,
            "name": "테스트 양식",
            "format": "docx",
            "created_at": "2025-01-01T00:00:00",
            "structure": {"sections": [], "fields": {}, "has_table": False, "raw_text_preview": ""},
        }
        mock_tmpl_redis_client.get.return_value = json.dumps(meta)

        response = tmpl_client.get(f"/api/v1/templates/{template_id}")
        assert response.status_code == 200

    def test_get_detail_has_structure_field(self, tmpl_client, mock_tmpl_redis_client):
        """상세 응답에 structure 포함"""
        template_id = str(uuid.uuid4())
        meta = {
            "template_id": template_id,
            "name": "테스트",
            "format": "pdf",
            "created_at": "2025-01-01T00:00:00",
            "structure": {
                "sections": [{"title": "개요", "level": 1}],
                "fields": {},
                "has_table": False,
                "raw_text_preview": "테스트",
            },
        }
        mock_tmpl_redis_client.get.return_value = json.dumps(meta)

        response = tmpl_client.get(f"/api/v1/templates/{template_id}")
        data = response.json()
        assert "structure" in data

    def test_get_detail_404_for_nonexistent(self, tmpl_client, mock_tmpl_redis_client):
        """존재하지 않는 template_id → 404"""
        mock_tmpl_redis_client.get.return_value = None

        response = tmpl_client.get(f"/api/v1/templates/{uuid.uuid4()}")
        assert response.status_code == 404

    def test_get_detail_returns_correct_template_id(self, tmpl_client, mock_tmpl_redis_client):
        """응답의 template_id가 요청한 ID와 일치"""
        template_id = str(uuid.uuid4())
        meta = {
            "template_id": template_id,
            "name": "테스트",
            "format": "docx",
            "created_at": "2025-01-01T00:00:00",
            "structure": {"sections": [], "fields": {}, "has_table": False, "raw_text_preview": ""},
        }
        mock_tmpl_redis_client.get.return_value = json.dumps(meta)

        response = tmpl_client.get(f"/api/v1/templates/{template_id}")
        data = response.json()
        assert data["template_id"] == template_id


# ---------------------------------------------------------------------------
# DELETE /api/v1/templates/{template_id} 테스트 (REQ-TMPL-003)
# ---------------------------------------------------------------------------


class TestDeleteTemplate:
    """DELETE /api/v1/templates/{template_id} - 템플릿 삭제"""

    def test_delete_existing_returns_204(self, tmpl_client, mock_tmpl_redis_client, tmp_path):
        """존재하는 템플릿 삭제 → 204"""
        template_id = str(uuid.uuid4())
        meta = {
            "template_id": template_id,
            "name": "삭제 테스트",
            "format": "docx",
            "created_at": "2025-01-01T00:00:00",
            "structure": {"sections": [], "fields": {}, "has_table": False, "raw_text_preview": ""},
        }
        # 디렉토리 생성
        tmpl_dir = tmp_path / "templates" / template_id
        tmpl_dir.mkdir(parents=True)
        (tmpl_dir / "original.docx").write_bytes(b"fake docx")

        mock_tmpl_redis_client.get.return_value = json.dumps(meta)

        response = tmpl_client.delete(f"/api/v1/templates/{template_id}")
        assert response.status_code == 204

    def test_delete_nonexistent_returns_404(self, tmpl_client, mock_tmpl_redis_client):
        """존재하지 않는 template_id 삭제 → 404"""
        mock_tmpl_redis_client.get.return_value = None

        response = tmpl_client.delete(f"/api/v1/templates/{uuid.uuid4()}")
        assert response.status_code == 404

    def test_delete_removes_redis_key(self, tmpl_client, mock_tmpl_redis_client, tmp_path):
        """삭제 후 Redis 키 제거 확인"""
        template_id = str(uuid.uuid4())
        meta = {
            "template_id": template_id,
            "name": "삭제 테스트",
            "format": "docx",
            "created_at": "2025-01-01T00:00:00",
            "structure": {"sections": [], "fields": {}, "has_table": False, "raw_text_preview": ""},
        }
        tmpl_dir = tmp_path / "templates" / template_id
        tmpl_dir.mkdir(parents=True)

        mock_tmpl_redis_client.get.return_value = json.dumps(meta)
        tmpl_client.delete(f"/api/v1/templates/{template_id}")

        # Redis delete가 호출되었는지 확인
        assert mock_tmpl_redis_client.delete.called
